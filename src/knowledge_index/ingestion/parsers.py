"""Concrete parsers for each supported MIME type (SPEC §7.1).

Heavy third-party libraries (pymupdf4llm, trafilatura, python-docx) are imported
lazily *inside* the parser that needs them, so importing this module — and the
markdown/plain-text paths that need no extra deps — works in any environment.
PDF tables are preserved as markdown (pymupdf4llm emits markdown natively).
"""

from __future__ import annotations

import hashlib
import re
from pathlib import Path

from common.errors import KnowledgeAgentError
from common.types import MimeType
from knowledge_index.ingestion.base import ParsedDoc, Parser, StructureElement
from knowledge_index.ingestion.normalize import decode_bytes, normalize_text

# --- mime detection --------------------------------------------------------

_EXT_TO_MIME = {
    ".pdf": MimeType.PDF,
    ".docx": MimeType.DOCX,
    ".html": MimeType.HTML,
    ".htm": MimeType.HTML,
    ".md": MimeType.MARKDOWN,
    ".markdown": MimeType.MARKDOWN,
    ".txt": MimeType.PLAIN,
    ".rst": MimeType.PLAIN,
    ".py": MimeType.CODE,
    ".js": MimeType.CODE,
    ".ts": MimeType.CODE,
    ".go": MimeType.CODE,
    ".java": MimeType.CODE,
    ".rs": MimeType.CODE,
}


def mime_from_path(path: str | Path) -> MimeType:
    """Map a file extension to a :class:`MimeType` (defaults to UNKNOWN)."""
    return _EXT_TO_MIME.get(Path(path).suffix.lower(), MimeType.UNKNOWN)


def _doc_id(blob: bytes) -> str:
    """Stable content-addressed doc id (sha256 prefix)."""
    return "doc-" + hashlib.sha256(blob).hexdigest()[:16]


# --- markdown --------------------------------------------------------------

_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
_FENCE = re.compile(r"^```")
_TABLE_ROW = re.compile(r"^\s*\|.*\|\s*$")
_LIST_ITEM = re.compile(r"^\s*([-*+]|\d+\.)\s+")


def _markdown_structure(text: str, title: str | None) -> list[StructureElement]:
    """Parse markdown into ordered structural elements."""
    elements: list[StructureElement] = []
    lines = text.split("\n")
    i = 0
    n = len(lines)
    para: list[str] = []

    def flush_para() -> None:
        if para:
            joined = " ".join(para).strip()
            if joined:
                elements.append(StructureElement(kind="paragraph", text=joined))
            para.clear()

    while i < n:
        line = lines[i]
        m = _HEADING.match(line)
        if m:
            flush_para()
            elements.append(
                StructureElement(kind="heading", text=m.group(2).strip(), level=len(m.group(1)))
            )
            i += 1
            continue
        if _FENCE.match(line):
            flush_para()
            block = [line]
            i += 1
            while i < n and not _FENCE.match(lines[i]):
                block.append(lines[i])
                i += 1
            if i < n:
                block.append(lines[i])  # closing fence
                i += 1
            elements.append(StructureElement(kind="code", text="\n".join(block)))
            continue
        if _TABLE_ROW.match(line):
            flush_para()
            block = [line]
            i += 1
            while i < n and _TABLE_ROW.match(lines[i]):
                block.append(lines[i])
                i += 1
            elements.append(StructureElement(kind="table", text="\n".join(block)))
            continue
        if _LIST_ITEM.match(line):
            flush_para()
            block = [line]
            i += 1
            while i < n and (_LIST_ITEM.match(lines[i]) or lines[i].startswith(("  ", "\t"))):
                block.append(lines[i])
                i += 1
            elements.append(StructureElement(kind="list", text="\n".join(block)))
            continue
        if line.strip() == "":
            flush_para()
            i += 1
            continue
        para.append(line)
        i += 1
    flush_para()
    return elements


class MarkdownParser:
    """Native markdown parser — no external dependencies."""

    async def parse(self, blob: bytes, hint: MimeType = MimeType.MARKDOWN) -> ParsedDoc:
        raw = decode_bytes(blob)
        text = normalize_text(raw)
        structure = _markdown_structure(text, None)
        title = next((e.text for e in structure if e.kind == "heading"), None)
        return ParsedDoc(
            doc_id=_doc_id(blob),
            text=text,
            structure=structure,
            metadata={"title": title, "parser": "markdown", "mime": str(hint)},
        )


class PlainTextParser:
    """Plain text / source-code parser: split on blank lines into paragraphs."""

    async def parse(self, blob: bytes, hint: MimeType = MimeType.PLAIN) -> ParsedDoc:
        text = normalize_text(decode_bytes(blob))
        kind = "code" if hint == MimeType.CODE else "paragraph"
        structure = [
            StructureElement(kind=kind, text=block.strip())  # type: ignore[arg-type]
            for block in re.split(r"\n\s*\n", text)
            if block.strip()
        ]
        return ParsedDoc(
            doc_id=_doc_id(blob),
            text=text,
            structure=structure,
            metadata={"parser": "plain", "mime": str(hint)},
        )


class HTMLParser:
    """HTML parser using trafilatura (lazy) with a regex fallback."""

    async def parse(self, blob: bytes, hint: MimeType = MimeType.HTML) -> ParsedDoc:
        raw = decode_bytes(blob)
        text: str | None = None
        title: str | None = None
        try:
            import trafilatura  # type: ignore

            extracted = trafilatura.extract(raw, output_format="markdown", include_tables=True)
            if extracted:
                text = extracted
            meta = trafilatura.extract_metadata(raw)
            if meta is not None:
                title = getattr(meta, "title", None)
        except ImportError:
            text = None
        if text is None:
            # Minimal fallback: strip tags, keep <title>.
            tm = re.search(r"<title>(.*?)</title>", raw, re.IGNORECASE | re.DOTALL)
            title = tm.group(1).strip() if tm else title
            stripped = re.sub(r"(?is)<(script|style).*?</\1>", " ", raw)
            stripped = re.sub(r"(?s)<[^>]+>", " ", stripped)
            text = stripped
        text = normalize_text(text)
        structure = _markdown_structure(text, title)
        return ParsedDoc(
            doc_id=_doc_id(blob),
            text=text,
            structure=structure,
            metadata={"title": title, "parser": "html", "mime": str(hint)},
        )


class PDFParser:
    """PDF parser using pymupdf4llm (markdown, tables preserved); requires the lib."""

    async def parse(self, blob: bytes, hint: MimeType = MimeType.PDF) -> ParsedDoc:
        try:
            import fitz  # type: ignore  # pymupdf
            import pymupdf4llm  # type: ignore
        except ImportError as e:  # pragma: no cover - depends on optional dep
            raise KnowledgeAgentError(
                "PDF parsing requires 'pymupdf4llm' (and 'pymupdf'); install the ingest extras"
            ) from e
        doc = fitz.open(stream=blob, filetype="pdf")
        try:
            md = pymupdf4llm.to_markdown(doc)
            title = (doc.metadata or {}).get("title") or None
            author = (doc.metadata or {}).get("author") or None
            page_count = doc.page_count
        finally:
            doc.close()
        text = normalize_text(md)
        structure = _markdown_structure(text, title)
        return ParsedDoc(
            doc_id=_doc_id(blob),
            text=text,
            structure=structure,
            metadata={
                "title": title,
                "author": author,
                "pages": page_count,
                "parser": "pymupdf4llm",
                "mime": str(hint),
            },
        )


class DocxParser:
    """DOCX parser using python-docx (lazy)."""

    async def parse(self, blob: bytes, hint: MimeType = MimeType.DOCX) -> ParsedDoc:
        try:
            import io

            import docx  # type: ignore  # python-docx
        except ImportError as e:  # pragma: no cover - depends on optional dep
            raise KnowledgeAgentError(
                "DOCX parsing requires 'python-docx'; install the ingest extras"
            ) from e
        document = docx.Document(io.BytesIO(blob))
        elements: list[StructureElement] = []
        parts: list[str] = []
        for para in document.paragraphs:
            t = para.text.strip()
            if not t:
                continue
            style = (para.style.name or "").lower() if para.style else ""
            if style.startswith("heading"):
                level = int(re.sub(r"\D", "", style) or "1")
                elements.append(StructureElement(kind="heading", text=t, level=min(level, 6)))
                parts.append("#" * min(level, 6) + " " + t)
            else:
                elements.append(StructureElement(kind="paragraph", text=t))
                parts.append(t)
        for table in document.tables:
            rows = [" | ".join(c.text.strip() for c in r.cells) for r in table.rows]
            md_table = "\n".join(f"| {r} |" for r in rows)
            elements.append(StructureElement(kind="table", text=md_table))
            parts.append(md_table)
        text = normalize_text("\n\n".join(parts))
        title = next((e.text for e in elements if e.kind == "heading"), None)
        props = document.core_properties
        return ParsedDoc(
            doc_id=_doc_id(blob),
            text=text,
            structure=elements,
            metadata={
                "title": title or (props.title or None),
                "author": props.author or None,
                "parser": "python-docx",
                "mime": str(hint),
            },
        )


_PARSERS: dict[MimeType, Parser] = {
    MimeType.MARKDOWN: MarkdownParser(),
    MimeType.HTML: HTMLParser(),
    MimeType.PDF: PDFParser(),
    MimeType.DOCX: DocxParser(),
    MimeType.PLAIN: PlainTextParser(),
    MimeType.CODE: PlainTextParser(),
    MimeType.UNKNOWN: PlainTextParser(),
}


def get_parser(hint: MimeType) -> Parser:
    """Return the registered parser for a MIME type (PlainText for UNKNOWN)."""
    return _PARSERS.get(hint, _PARSERS[MimeType.UNKNOWN])


async def parse_blob(blob: bytes, hint: MimeType) -> ParsedDoc:
    """Convenience: dispatch ``blob`` to the right parser by ``hint``."""
    return await get_parser(hint).parse(blob, hint)


async def parse_path(path: str | Path) -> ParsedDoc:
    """Read a file, detect its MIME from the extension, and parse it.

    The originating file path is recorded in ``metadata['path']``.
    """
    p = Path(path)
    blob = p.read_bytes()
    hint = mime_from_path(p)
    doc = await parse_blob(blob, hint)
    doc.metadata.setdefault("path", str(p))
    doc.metadata.setdefault("title", doc.metadata.get("title") or p.stem)
    return doc


__all__ = [
    "mime_from_path",
    "MarkdownParser",
    "PlainTextParser",
    "HTMLParser",
    "PDFParser",
    "DocxParser",
    "get_parser",
    "parse_blob",
    "parse_path",
]
